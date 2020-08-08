# -*- coding: utf-8 -*-
"""
Created on Sat Aug  8 11:27:08 2020

@author: Brenda Rojas Delgado
"""

import docx
import os
from docx.shared import Pt
from docx.shared import Inches

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.size = Pt(12)
main_header=doc.add_heading("CALCOLO NUMERICO PER LA FINANZA 87611",0)
doc.add_heading("IL DEBITO PUBBLICO SPAGNOLO",0)
doc.add_picture('C:/Users/15571272/Pictures/Saved Pictures/sampleimage.png',width=Inches(7),height=Inches(5))

nomi=doc.add_paragraph("")
nomi.add_run("Nomi: ").bold=True
doc.add_paragraph("Anastasiya Poyda Oleksyuk")
doc.add_paragraph("Numero di matricola: 1900081406")
doc.add_paragraph("Raúl Navarro Pedro")
doc.add_paragraph("Numero di matricola: 1900081405")

doc.add_heading("INTRODUZIONE",1)

intro1 = doc.add_paragraph("In questo documento, analizzeremo i risultati \
ottenuti in due periodi di tempo, l'emissione del debito pubblico spagnolo \
nel 2012 e nel 2018.")
intro1.alignment = 3
intro2 = doc.add_paragraph('Abbiamo scelto cinque titoli di ogni anno emessi \
dalla Segreteria Generale del Tesoro e della Finanza Internazionale, subordinata \
dalla Banca di Spagna attraverso aste e semplici operazioni di vendita del \
Tesoro Pubblico. Abbiamo basato la nostra selezione su obbligazioni e obbligazioni, \
che si caratterizzano come titoli finanziari con scadenza superiore a un anno \
(bollette). Il sistema di calcolo dei prezzi e dei rendimenti del debito sovrano \
è illustrato in un documento pubblicato dalla Banca di Spagna chiamato \
"Armonizzazione dei criteri di calcolo dei prezzi e dei rendimenti". criteri per \
il calcolo, la procedura, i termini, le tipologie di offerte, la presentazione e \
la risoluzione del debito pubblico sono riportati nella Gazzetta Ufficiale dello \
Stato all\'articolo 15 pubblicato ed entrato in vigore il 17/01/2020.')
intro2.alignment = 3
intro3 = doc.add_paragraph("Il rapporto si articola principalmente in quattro \
sezioni: in primo luogo, si inizia spiegando il contesto economico-finanziario \
del sistema spagnolo nei due periodi di tempo che abbiamo scelto; poi si analizzano \
i cinque dati sui titoli di Stato e sulle obbligazioni per ogni anno; in una terza \
sezione si presentano le formule che abbiamo calcolato e i rispettivi risultati; \
infine, in riferimento alle altre sezioni, si deduce una serie di idee rilevanti \
sull'argomento.")
intro3.alignment = 3

doc.add_heading("CONTESTO ECONOMICO E FINANZIARIO NEL PERIODO 2012 E 2018",1)

contesto1=doc.add_paragraph("In primo luogo, spiegheremo cos'è il ")
contesto1.add_run("debito pubblico ").bold=True
contesto1.add_run("e di quali attività è composto, e più tardi spiegheremo i due \
scenari per il settore economico spagnolo nel 2012 e nel 2018. ")
contesto1.alignment = 3

contesto2=doc.add_paragraph("Il debito dello Stato è costituito da capitale \
emesso dal governo attraverso aste pubbliche competitive, operazioni di credito, \
posizione debitoria di terzi e altri meccanismi di acquisto e vendita. L'obiettivo \
dell'emissione di titoli sovrani è quello di ottenere il finanziamento delle spese \
pubbliche e/o di ottenere liquidità. I principali titoli a reddito fisso sono \
le cambiali, le bond e le obbligazioni. Ora distingueremo i diversi tipi di titoli:")
contesto2.alignment = 3

contesto3=doc.add_paragraph("")
contesto3.add_run("Lettere dello tesoro").bold=True
contesto3.add_run(", la scadenza del titolo è a breve termine, inferiore o uguale \
a un anno.")
contesto3.alignment = 3

contesto4=doc.add_paragraph("")
contesto4.add_run("Bonos dello Estado").bold=True
contesto4.add_run(", la scadenza del titolo è a medio termine, durata da due a \
cinque anni. Attualmente lo Stato emette obbligazioni a tre e cinque anni.")
contesto4.alignment = 3

contesto5=doc.add_paragraph("")
contesto5.add_run("Obbligazioni governative").bold=True
contesto5.add_run(", la scadenza del titolo è a lungo termine, più di cinque anni. \
Attualmente il governo emette obbligazioni a dieci, quindici, trenta e cinquanta anni.")
contesto5.alignment = 3

doc.add_paragraph(
    'valore nominale = 1.000€', style='List Bullet 2'
)
doc.add_paragraph(
    'tali attività sono concordate a un tasso fisso di interesse pagato periodicamente \
sotto forma di cedole annuali post-pagabili (ad eccezione dei buoni del Tesoro che \
sono pagati in anticipo, prepagati)', style='List Bullet 2'
).alignment = 3
doc.add_paragraph(
    'emissione sul mercato primario e secondario, sopra la parità, alla pari, sotto \
la parità', style='List Bullet 2'
).alignment = 3
doc.add_paragraph(
    'il riscatto è alla pari = nominale di 1.000€', style='List Bullet 2'
)

doc.add_heading("SCENARIO ECONOMICO-FINANZIARIO NEL 2012",2)

doc.add_paragraph(
    'La crisi diffusa è iniziata nell\'estate del 2007, generando una crisi del \
debito pubblico e del deficit della bilancia dei pagamenti. Desidero sottolineare \
la necessità di introdurre cambiamenti e miglioramenti negli strumenti politici \
nella società spagnola, come in altre economie dell\'UEM. Le riforme sono state \
concepite per alleviare la recessione economico-finanziaria del paese.'
).alignment = 3

doc.add_paragraph(
    'L\'UEM si è basata sulle aspettative di una maggiore crescita potenziale all\'interno \
dell\'area, come ad esempio livelli più elevati di reddito pro capite in paesi \
con una crescita inferiore alla media europea. Ha inoltre facilitato l\'integrazione \
dei mercati commerciali e finanziari attraverso il consolidamento di una moneta \
unica e la libera circolazione dei flussi di capitali. Questi progressi hanno \
portato molti vantaggi ai paesi membri, ma hanno anche generato squilibri nella \
bilancia dei pagamenti e negli investimenti internazionali.'
).alignment = 3
    
doc.add_paragraph(
    'Il Trattato di Maastricht sull\'UE, entrato in vigore il 1° novembre 1993, \
comprende i tre pilastri fondamentali:'
).alignment = 3
    
doc.add_paragraph(
    'Comunità Europee→ accordi comuni su questioni e attività lavorative', 
    style='List Bullet 2'
).alignment = 3

doc.add_paragraph(
    'Cooperazione tra i membri dell’UE nel settore della giustizia e degli \
        affari interni→ controllo', 
    style='List Bullet 2'
).alignment = 3
    
doc.add_paragraph(
    'Cooperazione tra i membri dell’UE nel settore della giustizia e degli \
        affari interni→ controllo', 
    style='List Bullet 2'
).alignment = 3

doc.add_paragraph(
    'Vari organismi subordinano e correggono il debito in eccesso, fissando \
limiti a medio termine, mantenendo l\'inflazione vicino al 2% e non superando \
i limiti del debito sovrano. La clausola di no-bailout condiziona i governi a \
poter garantire il funzionamento dei mercati finanziari e a canalizzare i \
rischi di instabilità economica.', 
    style='List Bullet 2'
).alignment = 3
    
doc.add_paragraph(
    'Durante questo periodo, molti paesi europei hanno subito grandi perdite \
economiche e hanno accelerato il loro indebitamento senza che le misure di \
correzione economica basate sul trattato avessero un effetto sufficiente. \
In altre parole, non ha previsto meccanismi per attuare una risposta coordinata \
da parte dei membri dell\'UE agli squilibri della crisi del 2007.'    
).alignment = 3
    
doc.add_paragraph(
    'Nel breve termine, la BCE ha sostenuto finanziariamente i paesi per \
garantire la stabilità dell\'euro. All\'epoca la crisi era insufficiente, \
la crisi era di più ampia portata, gli strumenti di controllo e di \
coordinamento fiscale tra i paesi sono stati riconsiderati, sono stati \
introdotti nuovi meccanismi di vigilanza nel sistema finanziario e la gestione \
del rischio è stata resa più solidale. Il debito pubblico nel 2012 ha raggiunto \
l\'86,3% del PIL, generando un tasso di disoccupazione di 26%.'    
).alignment = 3
    
doc.add_paragraph(
    'Durante il periodo di unificazione della moneta unica, l\'euro (1999) \
fino al 2007, la crescita e i bassi costi di finanziamento hanno portato a forti \
squilibri fiscali nell\'economia spagnola (disavanzi sovrani). Il Governo ha \
approvato misure di riduzione delle imposte per ottenere entrate straordinarie \
e aumentare la spesa pubblica, mantenendo il saldo delle partite correnti in \
equilibrio. Ciò è avvenuto nel corso della prima fase della recessione, cioè \
all\'inizio della crisi, e si è protratto per anni.'    
).alignment = 3

doc.add_paragraph(
    'Dopo l\'inizio dello scenario di crisi economica, nel periodo 2007-2012, \
la situazione non è migliorata, ma al contrario è peggiorata. Il mercato degli \
investimenti Lehman Brothers ha annunciato il bancarotta, il governo spagnolo \
ha aumentato le garanzie in FDGD (Fondo di garanzia dei depositi) di 100.000 \
euro, avviando misure per ripulire il sistema finanziario. Per gli anni 2008-2011 \
è stato proposto un Patto di Stabilità che ha portato ad un aumento della \
disoccupazione, ad un aumento del deficit pubblico e ad un\'ulteriore recessione \
economica. Per alleviare la situazione finanziaria del paese sono stati effettuati \
raggruppamenti e fusioni, come ad esempio ristrutturazioni di banche (chiusure di \
casse di risparmio). Inoltre, sono state aumentate le aliquote fiscali, i salari \
sono stati ridotti e l\'età pensionabile è stata abbassata da 65 a 67 anni. \
Sebbene le misure siano state insufficienti, il tasso di disoccupazione nel 2012 \
ha raggiunto il 21,5%, il debito delle famiglie è passato dal 95% del 2007 al 100% \
del reddito lordo e il deficit pubblico è stato pari al 10,74% del PIL. Nel 2012 il \
sistema bancario spagnolo sarà seriamente discusso con i ministri dell\'economia e \
della competitività, chiedendo aiuto per ripulire il sistema.'    
).alignment = 3

doc.add_paragraph(
    'La politica di approvvigionamento della liquidità si è basata principalmente \
su operazioni d\'asta a tasso fisso con piena aggiudicazione, estensione delle \
garanzie collaterali nelle operazioni di prestito ed estensione fino a 3 anni \
dell\'orizzonte temporale delle sue operazioni a lungo termine (esempi di azioni \
della BCE); ciò ha permesso di alleviare le tensioni e di guadagnare il tempo \
necessario in modo che le azioni su altri fronti diano i loro frutti. È stata \
inoltre avviata la riforma delle politiche micro e macroprudenziali, per riparare \
gli errori che avevano consentito un eccessivo accumulo di rischio nella maggior \
parte dei sistemi bancari dell\'UEM. I paesi con una maggiore crescita del debito \
estero hanno registrato una minore crescita della produttività.'    
).alignment = 3

doc.add_heading("SCENARIO ECONOMICO-FINANZIARIO NEL 2018",2)

doc.add_paragraph(
    'Gli aspetti più caratteristici dello scenario economico per il 2018 sono \
stati la crescita e i mercati finanziari. Da un lato, il PIL è passato da un \
profilo di accelerazione all\'inizio dell\'anno a un profilo di moderata decelerazione. \
D\'altra parte, questo comportamento è stato molto diverso in diverse aree geografiche. \
Gli Stati Uniti avevano accelerato la crescita, trainata da una politica fiscale \
espansiva. La Cina ha adottato una gestione politica orientata ad una crescita \
economica moderata, un processo che sta gradualmente diventando più efficace. \
L\'Europa ha invece registrato un netto rallentamento influenzato da tre elementi: \
l\'aumento del prezzo del petrolio, l\'apprezzamento dell\'euro e le incertezze dei \
fattori macroeconomici.'    
).alignment = 3
    
doc.add_paragraph(
    'Durante la crisi, il tasso di inflazione dell\'UEM è stato inferiore all\'obiettivo \
(2%). Alcuni elementi di analisi ci permettono di spiegare la persistenza di questa \
situazione nel tempo: tassi di inflazione più bassi. Introduzione di riforme strutturali \
e di processi di aggiustamento competitivo e fiscale durante la crisi. Contribuendo così \
a una dinamica dei prezzi più contenuta, che avrebbe interessato settori ad alta intensità \
di manodopera come i servizi. Inoltre, la bassa sensibilità ciclica dell\'inflazione \
nell\'UEM spiega in parte perché l\'attuale fase di ripresa economica non è stata \
accompagnata da un aumento delle pressioni inflazionistiche. Inoltre, il contesto \
recessivo ha contribuito a ridurre le pressioni salariali, a causa del contenimento \
dei margini delle imprese (impatto inflazionistico).'    
).alignment = 3

doc.add_paragraph(
    'Fattori esplicativi a breve e medio termine (Spagna e UEM):'    
).alignment = 3
  
doc.add_paragraph(
    'Situazione recessiva (output gap negativo): basso tasso di risposta', 
    style='List Bullet 2'
).alignment = 3

doc.add_paragraph(
    'Prezzi dell\'energia (pressioni al ribasso, 2014-2017)', 
    style='List Bullet 2'
).alignment = 3
    
doc.add_paragraph(
    'Proroga della bassa inflazione: gli agenti abbassano le loro aspettative, \
non investono', 
    style='List Bullet 2'
).alignment = 3
    
doc.add_paragraph(
    'La situazione di bassa inflazione in questo momento è cruciale per il PM.\
In breve, nel 2017 e all\'inizio del 2018 l\'economia si stava riprendendo dalla \
recessione che durava da sette anni.'    
).alignment = 3
    
doc.add_heading("ANALISI DEI DATI SUL DEBITO PUBBLICO SPAGNOLO PER IL 2012 E IL 2018",1)

doc.add_paragraph(
    'In questa sezione spiegheremo i dati che abbiamo scelto per effettuare i calcoli \
appropriati.'    
).alignment = 3
    
doc.add_heading("DEBITO SOVRANO - 2012",2)
    
records1 = (
    ('ES0000011660 O EST 6.15 31.01.13',1,0.1,101.700,101.700,101.700,3.50,\
     '101,950 (30/05/2012)'),
         ('ES0000012866 O EST 4.20 30.07.13',1,5.30,100.010,100.010,100.010,4.18,'100.133 (30/05/2012)'
), ('ES0000012098 O EST 4.75 30.07.14',7,98.63,98.681,99.810,99.410,4.90,\
     '99,492 (30/05/2012) '
), ('ES0000012916 O EST 4.40 31.01.15',3,27.00,98.203,98.399,98.169,5.12,\
     '98,024 (30/05/2012)'
),
    ('ES00000120GA O EST 3.15 31.01.16',8,152.60,91.957,92.100,91.000,5.64,\
     '92,700 (29/05/2012)'
),
)

table1 = doc.add_table(rows=1, cols=8)
caption=doc.add_paragraph("")
caption.add_run("Tavolo 1").bold=True
caption.add_run(": Emissione del debito pubblico spagnolo il 31/05/2012.")
footnote=doc.add_paragraph("Nota: gli importi sono in milioni di € e i prezzi sono in %")
caption.alignment = 1
table1.style = 'TableGrid'
table1.autofit=True
hdr_cells = table1.rows[0].cells
em='EMISSIONE'
op='Nº DI OPERAZIONI'
im='IMPORTO CONTRATTUALE'
pr_mz='PREZZO (EX-COUPON) MEZZO'
pr_mx='PREZZO (EX-COUPON) MASSIMO'
pr_mn='PREZZO (EX-COUPON) MINIMO'
rn_mz='RENDTO. INTERNO MEZZO'
pr_mz_pd='PREZZO MEZZO PRECEDENTE (DATA)'
hdr_cells[0].text = em
hdr_cells[1].text = op
hdr_cells[2].text = im
hdr_cells[3].text = pr_mz
hdr_cells[4].text = pr_mx
hdr_cells[5].text = pr_mn
hdr_cells[6].text = rn_mz
hdr_cells[7].text =pr_mz_pd
for em,op,im,pr_mz,pr_mx,pr_mn,rn_mz,pr_mz_pd in records1:
    row_cells = table1.add_row().cells
    row_cells[0].text = str(em)
    row_cells[1].text = str(op)
    row_cells[2].text = str(im)
    row_cells[3].text = str(pr_mz)
    row_cells[4].text = str(pr_mx)
    row_cells[5].text = str(pr_mn)
    row_cells[6].text = str(rn_mz)
    row_cells[7].text = str(pr_mz_pd)
    
for row in table1.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(8)
                
para1 = doc.add_paragraph("Nella tabella qui sopra, troviamo cinque tipi di \
legami con caratteristiche simili, se guardiamo il comportamento dei legami, \
vediamo che è il contrario. Più alto è il rendimento interno, più basso è il \
prezzo d'acquisto e l'interesse della cedola. Il primo tipo di obbligazione, \
ad esempio, è un titolo di Stato emesso il 31/05/2012. Ha un rendimento esplicito \
del 3,5% con una cedola annuale del 6,15%. L'asta è stata aggiudicata a un prezzo \
medio del 101,95% ('1001,95), il che significa che gli acquirenti hanno acquistato \
debiti a lungo termine a quel prezzo, con scadenza 31/01/2013.")
para1.alignment = 3

para2 = doc.add_paragraph("Da notare che i tassi di interesse nel 2012 sull'emissione \
di operazioni definitive da parte dello Stato spagnolo sono positivi, raggiungendo \
anche più di 5 punti. Si tratta di un dato molto significativo prima della crisi \
internazionale, in quanto questo tasso di interesse ha subito un drastico cambiamento \
di tendenza, che analizzeremo in seguito.")
para2.alignment = 3

para3 = doc.add_paragraph("Abbiamo scoperto che più lunga è la scadenza (durata \
dell'attività), più gli investitori sono disposti ad acquistare, in quanto il prezzo \
di acquisto è inferiore, ha un rendimento annuo inferiore ma a più lungo termine. \
Questi investitori cercano di investire in obbligazioni più rischiose ma con un \
rendimento futuro più elevato.")
para3.alignment = 3

p1 = para1._p
p1.addnext(para2._p)
p2=para2._p
p2.addnext(para3._p)          

doc.add_page_break()

doc.add_heading("DEBITO SOVRANO - 2018",2)   

records2 = (
    ('ES00000127D6 B EST 0.25 30.04.18','4','93.12','100.033','100.040','100.028','-0.56',\
     '100,027 (12/04/2018)'),
('ES00000124B7 B EST 3.75 31.10.18','2','20.00','102.294','102.300','102.288','-0.47',\
 '102,282 (12/04/2018)'),
('ES00000128A0 B EST O.25 31.01.19','2','198.81','100.547','100.550','100.544','-0.43',\
'100,546 (12/04/2018)'),
('ES00000122T3 0 EST 4.85 31.10.20','4','82.00','112.943','112.950','112.935','-0.22',\
'112,908 (12/04/2018) '),
('ES00000123B9 0 EST 5.50 30.04.21','8','200.44','117.917','117.212','117.184','-0.15',\
 '117,157 (12/04/2018)'),
)

table2 = doc.add_table(rows=1, cols=8)
caption=doc.add_paragraph("")
caption.add_run("Tavolo 2").bold=True
caption.add_run(": Emissione del debito pubblico spagnolo il 12/04/2018.")
footnote=doc.add_paragraph("Nota: gli importi sono in milioni di € e i prezzi sono in %")
caption.alignment = 1
table2.style = 'TableGrid'
table2.autofit=True
hdr_cells = table2.rows[0].cells
em='EMISSIONE'
op='Nº DI OPERAZIONI'
im='IMPORTO CONTRATTUALE'
pr_mz='PREZZO (EX-COUPON) MEZZO'
pr_mx='PREZZO (EX-COUPON) MASSIMO'
pr_mn='PREZZO (EX-COUPON) MINIMO'
rn_mz='RENDTO. INTERNO MEZZO'
pr_mz_pd='PREZZO MEZZO PRECEDENTE (DATA)'
hdr_cells[0].text = em
hdr_cells[1].text = op
hdr_cells[2].text = im
hdr_cells[3].text = pr_mz
hdr_cells[4].text = pr_mx
hdr_cells[5].text = pr_mn
hdr_cells[6].text = rn_mz
hdr_cells[7].text =pr_mz_pd
for em,op,im,pr_mz,pr_mx,pr_mn,rn_mz,pr_mz_pd in records2:
    row_cells = table2.add_row().cells
    row_cells[0].text = str(em)
    row_cells[1].text = str(op)
    row_cells[2].text = str(im)
    row_cells[3].text = str(pr_mz)
    row_cells[4].text = str(pr_mx)
    row_cells[5].text = str(pr_mn)
    row_cells[6].text = str(rn_mz)
    row_cells[7].text = str(pr_mz_pd)
    
for row in table2.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.size= Pt(8)  
                
para4 = doc.add_paragraph("Siamo nell'anno 2018, che è un anno importante per \
noi per valutare le aste del debito sovrano perché, nonostante la crescita \
economica, gli Stati sono fortemente indebitati e la politica della BCE di tassi \
di interesse molto bassi, vicino allo 0% o anche negativo, rende il debito degli \
Stati non una buona opportunità o una risorsa interessante per gli investitori \
che vogliono fare soldi perché i rendimenti sono negativi. Pur avendo un tasso \
di interesse positivo (ma basso), tenendo conto del prezzo pagato per la cedola \
e del periodo di tempo, il rendimento è negativo.")
para4.alignment = 3

para5 = doc.add_paragraph("Nella tabella, vediamo il debito a breve come i primi \
tre coupon e il debito a lungo termine (più di un anno) come potrebbero essere i \
restanti due coupon. Il periodo di tempo rimanente per finalizzare la cedola è \
importante, poiché vediamo come il periodo è più lungo, le cedole hanno un rendimento \
negativo inferiore e, quindi, un rendimento migliore per l'investitore.")
para5.alignment = 3

para6 = doc.add_paragraph("Questo debito pubblico, nonostante abbia rendimenti \
negativi, è scelto da alcuni investitori o istituzioni perché hanno un rendimento \
migliore rispetto al possesso di denaro nei depositi bancari o, ad esempio, per \
venderlo in futuro e realizzare un profitto.")
para6.alignment = 3

p4 = para4._p
p4.addnext(para5._p)
p5=para5._p
p5.addnext(para6._p)
 
doc.save("debito_pubblico.docx")
os.system("start debito_pubblico.docx")